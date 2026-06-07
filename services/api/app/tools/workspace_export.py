from __future__ import annotations

import hashlib
import json
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from app.tools.audit_log import append_audit_event, read_audit_log, verify_audit_hash_chain
from app.tools.bibtex_generator import read_bibtex
from app.tools.citation_grounding import read_citation_grounding_report
from app.tools.file_tools import ensure_dir, relative_posix, write_json, write_text
from app.tools.literature_index import load_literature_index
from app.tools.manuscript_references import read_references_status
from app.tools.readiness_report import generate_v1_readiness_report
from app.tools.statistical_assistant import generate_statistical_assistant_report


WORKSPACE_EXPORT_DIR = "exports/workspace"
DOCX_PATH = f"{WORKSPACE_EXPORT_DIR}/research_workspace_export.docx"
LATEX_PATH = f"{WORKSPACE_EXPORT_DIR}/research_workspace_export.tex"
TRUST_MD_PATH = f"{WORKSPACE_EXPORT_DIR}/trust_report.md"
TRUST_JSON_PATH = f"{WORKSPACE_EXPORT_DIR}/trust_report.json"
MANIFEST_PATH = f"{WORKSPACE_EXPORT_DIR}/workspace_export_manifest.json"

TEXT_SUFFIXES = {".json", ".md", ".tex"}
SECRET_PATTERNS = [
    re.compile(r"sk_live_[A-Za-z0-9_-]+"),
    re.compile(r"sk-[A-Za-z0-9_-]{20,}"),
    re.compile(r"OPENAI_API_KEY\s*[:=]", re.IGNORECASE),
    re.compile(r"BEGIN (?:RSA |EC |OPENSSH |)PRIVATE KEY"),
    re.compile(
        r"(api[_-]?key|secret|token|password)\s*[:=]\s*[\"']?[^\"'\s,;]{8,}",
        re.IGNORECASE,
    ),
]
ABSOLUTE_PATH_PATTERNS = [
    re.compile(r"[A-Za-z]:[\\/][^\s\"']+"),
    re.compile(r"/(?:home|Users|var|tmp|mnt)/[^\s\"']+"),
]
LATEX_ESCAPE_MAP = {
    "\\": r"\textbackslash{}",
    "&": r"\&",
    "%": r"\%",
    "$": r"\$",
    "#": r"\#",
    "_": r"\_",
    "{": r"\{",
    "}": r"\}",
    "~": r"\textasciitilde{}",
    "^": r"\textasciicircum{}",
}


class WorkspaceExportError(ValueError):
    pass


def _utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _project_path(project_dir: Path, relative_path: str) -> Path:
    root = project_dir.resolve()
    path = (root / relative_path).resolve()
    try:
        path.relative_to(root)
    except ValueError as exc:
        raise WorkspaceExportError(f"Unsafe workspace export path: {relative_path}") from exc
    return path


def _read_json(project_dir: Path, relative_path: str, fallback: Any) -> Any:
    path = _project_path(project_dir, relative_path)
    if not path.exists():
        return fallback
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return fallback


def _read_text(project_dir: Path, relative_path: str) -> str:
    path = _project_path(project_dir, relative_path)
    if not path.exists():
        return ""
    return path.read_text(encoding="utf-8", errors="replace")


def _shorten(value: Any, limit: int = 700) -> str:
    text = str(value or "").strip()
    text = re.sub(r"\s+", " ", text)
    if len(text) <= limit:
        return text
    return text[: limit - 3].rstrip() + "..."


def _latex_escape(value: Any) -> str:
    return "".join(LATEX_ESCAPE_MAP.get(char, char) for char in str(value or ""))


def _relative_artifact(path: Path, project_dir: Path) -> str:
    return relative_posix(path, project_dir)


def _scan_generated_text(project_dir: Path, relative_path: str) -> list[str]:
    path = _project_path(project_dir, relative_path)
    if not path.exists() or path.suffix.lower() not in TEXT_SUFFIXES:
        return []
    text = path.read_text(encoding="utf-8", errors="replace")
    warnings: list[str] = []
    for pattern in SECRET_PATTERNS:
        if pattern.search(text):
            warnings.append(f"{relative_path}: secret-like pattern")
            break
    for pattern in ABSOLUTE_PATH_PATTERNS:
        if pattern.search(text):
            warnings.append(f"{relative_path}: absolute-path-like pattern")
            break
    return warnings


def _artifact_entry(
    project_dir: Path,
    relative_path: str,
    artifact_type: str,
    mime_type: str,
    required: bool = True,
) -> dict[str, Any]:
    path = _project_path(project_dir, relative_path)
    if not path.exists():
        return {
            "artifact_type": artifact_type,
            "relative_path": relative_path,
            "mime_type": mime_type,
            "available": False,
            "required": required,
            "size_bytes": 0,
            "sha256": None,
        }
    return {
        "artifact_type": artifact_type,
        "relative_path": relative_path,
        "mime_type": mime_type,
        "available": True,
        "required": required,
        "size_bytes": path.stat().st_size,
        "sha256": _sha256(path),
    }


def _source_file_entry(project_dir: Path, relative_path: str) -> dict[str, Any]:
    path = _project_path(project_dir, relative_path)
    return {
        "relative_path": relative_path,
        "available": path.exists() and path.is_file(),
        "size_bytes": path.stat().st_size if path.exists() and path.is_file() else 0,
    }


def _references_overview(references_status: dict[str, Any]) -> dict[str, Any]:
    verified = references_status.get("verified_references", [])
    candidates = references_status.get("candidate_references", [])
    placeholders = references_status.get("placeholder_records", [])
    return {
        "verified_references": len(verified) if isinstance(verified, list) else 0,
        "candidate_references": len(candidates) if isinstance(candidates, list) else 0,
        "placeholder_records": len(placeholders) if isinstance(placeholders, list) else 0,
        "warnings": references_status.get("warnings", [])
        if isinstance(references_status.get("warnings"), list)
        else [],
    }


def _build_trust_report(
    project_dir: Path,
    project_id: str,
    generated_at: str,
    readiness: dict[str, Any],
    trust: dict[str, Any],
    references_status: dict[str, Any],
    citation_grounding: dict[str, Any],
    statistical_report: dict[str, Any] | None,
) -> dict[str, Any]:
    audit_hash = verify_audit_hash_chain(project_dir)
    audit_entries = read_audit_log(project_dir, limit=0)
    literature = load_literature_index(project_dir)
    evidence_claims = _read_json(project_dir, "provenance/evidence.json", [])
    bibtex_report = read_bibtex(project_dir, project_id).get("report", {})
    references = _references_overview(references_status)

    return {
        "report_id": "workspace_trust_report_v15",
        "project_id": project_id,
        "generated_at": generated_at,
        "relative_path": TRUST_JSON_PATH,
        "scope": "local_mvp_workspace_export",
        "trust_summary": {
            "overall_status": trust.get("overall_status"),
            "scores": trust.get("scores", {}),
            "counts": trust.get("counts", {}),
            "open_items": trust.get("open_items", []),
            "blocking_issues": trust.get("blocking_issues", []),
        },
        "readiness": {
            "readiness_level": readiness.get("readiness_level"),
            "local_mvp_checks": readiness.get("local_mvp_checks", {}),
            "blocking_gaps": readiness.get("blocking_gaps", []),
            "known_gaps": readiness.get("production_gaps", []),
        },
        "references": references,
        "bibtex": {
            "formal_entries": bibtex_report.get("formal_entries", 0),
            "approved_entries": bibtex_report.get("approved_entries", 0),
            "candidate_records": bibtex_report.get("candidate_records", 0),
            "placeholder_records": bibtex_report.get("placeholder_records", 0),
            "warnings": bibtex_report.get("warnings", []),
        },
        "citation_grounding": {
            "summary": citation_grounding.get("summary", {}),
            "limitations": sorted(
                {
                    str(limit)
                    for item in citation_grounding.get("items", [])
                    if isinstance(item, dict)
                    for limit in item.get("limitations", [])
                }
            ),
        },
        "analysis": {
            "statistical_assistant_report": statistical_report.get("relative_path")
            if statistical_report
            else None,
            "statistical_assistant_available": statistical_report is not None,
            "guardrails": statistical_report.get("guardrails", []) if statistical_report else [],
        },
        "audit": {
            "hash_chain": audit_hash,
            "entry_count": len(audit_entries),
        },
        "source_counts": {
            "literature_records": len(literature),
            "evidence_claims": len(evidence_claims) if isinstance(evidence_claims, list) else 0,
        },
        "source_files": {
            "manuscript_draft": "manuscript/draft.md",
            "literature_index": "literature/literature_index.json",
            "evidence": "provenance/evidence.json",
            "trust_summary": "trust/trust_summary.json",
            "readiness_report": "trust/v1_readiness_report.json",
            "references_status": "manuscript/references_status.json",
            "citation_grounding": "provenance/citation_grounding_report.json",
            "bibtex_report": "literature/bibtex_report.json",
            "audit_log": "audit/audit_log.jsonl",
        },
        "caveats": [
            "This report summarizes local ResearchAgent artifacts only.",
            "It is not a production compliance archive, peer review certificate, plagiarism report, DOI verification, or scientific validation.",
            "References are not promoted unless the local approval workflow has been applied.",
            "No inferential statistics, significance statements, or causal conclusions are created by this export.",
        ],
    }


def _trust_report_markdown(report: dict[str, Any]) -> str:
    references = report["references"]
    readiness = report["readiness"]
    trust = report["trust_summary"]
    citation = report["citation_grounding"]
    audit = report["audit"]["hash_chain"]
    lines = [
        "# ResearchAgent Workspace Trust Report",
        "",
        f"- Project ID: `{report['project_id']}`",
        f"- Generated at: `{report['generated_at']}`",
        f"- Scope: `{report['scope']}`",
        "",
        "This report summarizes local workflow artifacts only. It is not a production compliance archive, peer review certificate, plagiarism report, DOI verification, or scientific validation.",
        "",
        "## Status",
        "",
        f"- Trust summary: `{trust.get('overall_status')}`",
        f"- Readiness level: `{readiness.get('readiness_level')}`",
        f"- Audit hash chain valid: `{audit.get('valid')}`",
        f"- Audit entries checked: `{audit.get('checked_entries')}`",
        "",
        "## References",
        "",
        f"- Approved verified references: `{references['verified_references']}`",
        f"- Candidate references: `{references['candidate_references']}`",
        f"- Placeholder or unverified records: `{references['placeholder_records']}`",
        "",
        "Formal references and BibTeX entries are limited to records that have verified metadata, human verification, and an applied approval decision.",
        "",
        "## Citation Grounding",
        "",
    ]
    for key, value in sorted((citation.get("summary") or {}).items()):
        lines.append(f"- `{key}`: `{value}`")
    lines.extend(
        [
            "",
            "## Blocking Items",
            "",
        ]
    )
    blocking = trust.get("blocking_issues", [])
    if isinstance(blocking, list) and blocking:
        for item in blocking:
            if isinstance(item, dict):
                lines.append(
                    f"- `{item.get('item_type')}` `{item.get('item_id')}`: {item.get('message')}"
                )
    else:
        lines.append("- No blocking item was returned by the local trust summary.")
    lines.extend(["", "## Known Gaps", ""])
    for gap in readiness.get("known_gaps", []):
        lines.append(f"- {gap}")
    lines.extend(["", "## Caveats", ""])
    for caveat in report.get("caveats", []):
        lines.append(f"- {caveat}")
    return "\n".join(lines).rstrip() + "\n"


def _latex_reference_lines(references_status: dict[str, Any]) -> list[str]:
    verified = references_status.get("verified_references", [])
    if not isinstance(verified, list) or not verified:
        return [
            "No approved verified references are available for formal bibliography output."
        ]
    lines: list[str] = []
    for entry in verified:
        if not isinstance(entry, dict):
            continue
        authors = ", ".join(str(author) for author in entry.get("authors", []) if str(author).strip())
        parts = [part for part in [authors, entry.get("year"), entry.get("title"), entry.get("journal")] if part]
        if entry.get("doi"):
            parts.append(f"DOI: {entry['doi']}")
        lines.append(". ".join(str(part) for part in parts))
    return lines


def _latex_document(
    report: dict[str, Any],
    references_status: dict[str, Any],
    manuscript_excerpt: str,
) -> str:
    references = _latex_reference_lines(references_status)
    citation_summary = report["citation_grounding"].get("summary", {})
    lines = [
        r"\documentclass[11pt]{article}",
        r"\usepackage[margin=1in]{geometry}",
        r"\usepackage[T1]{fontenc}",
        r"\usepackage[utf8]{inputenc}",
        r"\usepackage{longtable}",
        r"\title{ResearchAgent Workspace Export}",
        rf"\date{{{_latex_escape(report['generated_at'])}}}",
        r"\begin{document}",
        r"\maketitle",
        r"\section*{Local Scope}",
        _latex_escape(
            "This LaTeX source is generated from local ResearchAgent artifacts. It is not a production compliance archive, peer review certificate, plagiarism report, DOI verification, or scientific validation."
        ),
        r"\section*{Trust Summary}",
        r"\begin{itemize}",
        rf"\item Project ID: \texttt{{{_latex_escape(report['project_id'])}}}",
        rf"\item Trust status: \texttt{{{_latex_escape(report['trust_summary'].get('overall_status'))}}}",
        rf"\item Readiness level: \texttt{{{_latex_escape(report['readiness'].get('readiness_level'))}}}",
        rf"\item Audit hash chain valid: \texttt{{{_latex_escape(report['audit']['hash_chain'].get('valid'))}}}",
        r"\end{itemize}",
        r"\section*{Citation Grounding Summary}",
        r"\begin{itemize}",
    ]
    for key, value in sorted(citation_summary.items()):
        lines.append(rf"\item {_latex_escape(key)}: \texttt{{{_latex_escape(value)}}}")
    lines.extend(
        [
            r"\end{itemize}",
            r"\section*{Manuscript Excerpt}",
            _latex_escape(manuscript_excerpt or "No manuscript draft was found."),
            r"\section*{Approved References}",
            r"\begin{enumerate}",
        ]
    )
    for reference in references:
        lines.append(rf"\item {_latex_escape(reference)}")
    lines.extend(
        [
            r"\end{enumerate}",
            r"\section*{Export Caveats}",
            r"\begin{itemize}",
        ]
    )
    for caveat in report.get("caveats", []):
        lines.append(rf"\item {_latex_escape(caveat)}")
    lines.extend([r"\end{itemize}", r"\end{document}", ""])
    return "\n".join(lines)


def _write_docx(
    project_dir: Path,
    report: dict[str, Any],
    references_status: dict[str, Any],
    manuscript_excerpt: str,
) -> None:
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError as exc:
        raise WorkspaceExportError("python-docx is required to create the workspace DOCX export") from exc

    document = Document()
    for section in document.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    document.add_heading("ResearchAgent Workspace Export", level=0)
    document.add_paragraph(f"Project ID: {report['project_id']}")
    document.add_paragraph(f"Generated at: {report['generated_at']}")
    document.add_paragraph(
        "This document summarizes local ResearchAgent artifacts only. It is not a production compliance archive, peer review certificate, plagiarism report, DOI verification, or scientific validation."
    )

    document.add_heading("Trust Summary", level=1)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Signal"
    table.rows[0].cells[1].text = "Value"
    rows = [
        ("Trust status", report["trust_summary"].get("overall_status")),
        ("Readiness level", report["readiness"].get("readiness_level")),
        ("Audit hash chain valid", report["audit"]["hash_chain"].get("valid")),
        ("Approved verified references", report["references"]["verified_references"]),
        ("Candidate references", report["references"]["candidate_references"]),
        ("Placeholder records", report["references"]["placeholder_records"]),
    ]
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(value)

    document.add_heading("Manuscript Excerpt", level=1)
    document.add_paragraph(manuscript_excerpt or "No manuscript draft was found.")

    document.add_heading("Approved References", level=1)
    for reference in _latex_reference_lines(references_status):
        document.add_paragraph(reference, style="List Number")

    document.add_heading("Caveats", level=1)
    for caveat in report.get("caveats", []):
        document.add_paragraph(caveat, style="List Bullet")

    document.save(_project_path(project_dir, DOCX_PATH))


def _write_manifest(
    project_dir: Path,
    project_id: str,
    generated_at: str,
    warnings: list[str],
) -> dict[str, Any]:
    artifacts = [
        _artifact_entry(
            project_dir,
            DOCX_PATH,
            "word_docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ),
        _artifact_entry(project_dir, LATEX_PATH, "latex_source", "application/x-tex"),
        _artifact_entry(project_dir, TRUST_MD_PATH, "trust_report_markdown", "text/markdown"),
        _artifact_entry(project_dir, TRUST_JSON_PATH, "trust_report_json", "application/json"),
    ]
    source_files = [
        _source_file_entry(project_dir, relative_path)
        for relative_path in [
            "manuscript/draft.md",
            "literature/literature_index.json",
            "provenance/evidence.json",
            "trust/trust_summary.json",
            "trust/v1_readiness_report.json",
            "manuscript/references_status.json",
            "provenance/citation_grounding_report.json",
            "literature/bibtex_report.json",
            "analysis/statistical_assistant_report.json",
            "audit/audit_log.jsonl",
        ]
    ]
    scan_warnings = []
    for artifact in artifacts:
        if artifact["available"]:
            scan_warnings.extend(_scan_generated_text(project_dir, artifact["relative_path"]))
    all_warnings = list(dict.fromkeys(warnings + scan_warnings))
    manifest = {
        "available": True,
        "export_id": "workspace_export_v15",
        "project_id": project_id,
        "generated_at": generated_at,
        "relative_path": MANIFEST_PATH,
        "export_dir": WORKSPACE_EXPORT_DIR,
        "artifacts": artifacts
        + [
            {
                "artifact_type": "workspace_export_manifest",
                "relative_path": MANIFEST_PATH,
                "mime_type": "application/json",
                "available": True,
                "required": True,
                "size_bytes": 0,
                "sha256": None,
            }
        ],
        "source_files": source_files,
        "safety": {
            "project_relative_paths_only": True,
            "secret_scan_passed": not scan_warnings,
            "warning_count": len(all_warnings),
        },
        "warnings": all_warnings,
        "caveats": [
            "Workspace export is a local MVP artifact package, not a production backup.",
            "Generated DOCX and LaTeX are drafts for human review.",
            "The trust report is a local workflow summary, not scientific or compliance validation.",
        ],
    }
    write_json(_project_path(project_dir, MANIFEST_PATH), manifest)
    manifest_path = _project_path(project_dir, MANIFEST_PATH)
    manifest["artifacts"][-1]["size_bytes"] = manifest_path.stat().st_size
    write_json(manifest_path, manifest)
    return manifest


def build_workspace_export(project_dir: Path, project_id: str) -> dict[str, Any]:
    project_dir = project_dir.resolve()
    if not project_dir.exists() or not project_dir.is_dir():
        raise FileNotFoundError(f"Project directory does not exist: {project_id}")
    export_dir = ensure_dir(_project_path(project_dir, WORKSPACE_EXPORT_DIR))
    if project_dir not in export_dir.parents:
        raise WorkspaceExportError("Workspace export path escaped the project directory.")

    generated_at = _utc_now()
    readiness = generate_v1_readiness_report(project_dir, project_id)
    trust = _read_json(project_dir, "trust/trust_summary.json", {})
    references_status = read_references_status(project_dir, project_id)
    citation_grounding = read_citation_grounding_report(project_dir, project_id)
    statistical_report: dict[str, Any] | None = None
    statistical_warnings: list[str] = []
    try:
        statistical_report = generate_statistical_assistant_report(project_dir, project_id)
    except FileNotFoundError:
        statistical_warnings.append("analysis/statistical_assistant_report.json was not generated because analysis source files are missing.")

    manuscript_excerpt = _shorten(_read_text(project_dir, "manuscript/draft.md"), limit=2000)
    report = _build_trust_report(
        project_dir,
        project_id,
        generated_at,
        readiness,
        trust,
        references_status,
        citation_grounding,
        statistical_report,
    )

    write_json(_project_path(project_dir, TRUST_JSON_PATH), report)
    write_text(_project_path(project_dir, TRUST_MD_PATH), _trust_report_markdown(report))
    write_text(
        _project_path(project_dir, LATEX_PATH),
        _latex_document(report, references_status, manuscript_excerpt),
    )
    _write_docx(project_dir, report, references_status, manuscript_excerpt)
    manifest = _write_manifest(project_dir, project_id, generated_at, statistical_warnings)

    append_audit_event(
        project_dir,
        project_id,
        "generate_workspace_export",
        "Workspace DOCX, LaTeX, trust report, and manifest were generated from local artifacts.",
        {
            "manifest_file": MANIFEST_PATH,
            "docx_file": DOCX_PATH,
            "latex_file": LATEX_PATH,
            "trust_report_json": TRUST_JSON_PATH,
            "trust_report_markdown": TRUST_MD_PATH,
            "artifact_count": len(manifest["artifacts"]),
        },
        source="api",
        event_category="audit",
        risk_level="low",
        entity_type="audit_export",
        entity_id="workspace_export_v15",
    )
    return manifest


def latest_workspace_export_info(project_dir: Path, project_id: str) -> dict[str, Any]:
    project_dir = project_dir.resolve()
    manifest_path = _project_path(project_dir, MANIFEST_PATH)
    if not manifest_path.exists():
        return {
            "available": False,
            "export_id": "workspace_export_v15",
            "project_id": project_id,
            "relative_path": None,
            "export_dir": WORKSPACE_EXPORT_DIR,
            "artifacts": [],
            "source_files": [],
            "safety": {
                "project_relative_paths_only": True,
                "secret_scan_passed": True,
                "warning_count": 0,
            },
            "warnings": [],
            "caveats": [
                "No workspace export has been generated yet.",
            ],
            "message": "No workspace export has been generated yet.",
        }
    payload = json.loads(manifest_path.read_text(encoding="utf-8"))
    if not isinstance(payload, dict):
        raise WorkspaceExportError("Workspace export manifest is invalid.")
    payload["available"] = True
    return payload
