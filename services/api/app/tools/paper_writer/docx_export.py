from __future__ import annotations

import hashlib
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document

from app.tools.audit_log import append_audit_event
from app.tools.auto_scientist.scientist_paper import AUTONOMOUS_PAPER_MD
from app.tools.file_tools import ensure_dir, write_json
from app.tools.paper_writer.section_writer import DRAFT_FULL_MD

DRAFT_FULL_DOCX = "manuscript/draft_full.docx"
DRAFT_FULL_DOCX_MANIFEST = "manuscript/draft_full_docx_export_manifest.json"
AUTO_SCIENTIST_PAPER_DOCX = "manuscript/auto_scientist_paper.docx"
AUTO_SCIENTIST_PAPER_DOCX_MANIFEST = "manuscript/auto_scientist_paper_docx_export_manifest.json"

DOCX_DRAFT_CAVEAT = (
    "This DOCX is an AI-generated draft artifact for human review. It is not peer review, "
    "citation verification, scientific proof, or publication readiness."
)

SECRET_PATTERNS = [
    re.compile(r"sk_live_[A-Za-z0-9_-]+"),
    re.compile(r"sk-[A-Za-z0-9_-]{20,}"),
    re.compile(r"OPENAI_API_KEY\s*[:=]", re.IGNORECASE),
    re.compile(r"BEGIN (?:RSA |EC |OPENSSH |)PRIVATE KEY"),
    re.compile(
        r"(api[_-]?key|secret|token|password)\s*[:=]\s*[\"']?[^\"'\s,;]{8,}",
        re.IGNORECASE,
    ),
]
ABSOLUTE_PATH_PATTERNS = [
    re.compile(r"[A-Za-z]:[\\/][^\s\"')\]]+"),
    re.compile(r"/(?:home|Users|var|tmp|mnt)/[^\s\"')\]]+"),
]


class PaperDocxExportError(ValueError):
    pass


def _utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _project_path(project_dir: Path, relative_path: str) -> Path:
    root = project_dir.resolve()
    normalized = relative_path.replace("\\", "/")
    path = (root / normalized).resolve()
    try:
        path.relative_to(root)
    except ValueError as exc:
        raise PaperDocxExportError(f"Unsafe paper DOCX export path: {relative_path}") from exc
    return path


def _read_source_markdown(project_dir: Path, relative_path: str) -> str:
    path = _project_path(project_dir, relative_path)
    if not path.exists():
        raise FileNotFoundError(f"{relative_path} does not exist; generate the manuscript draft first")
    return path.read_text(encoding="utf-8", errors="replace")


def _validate_manuscript_markdown_path(relative_path: str) -> str:
    normalized = relative_path.strip().replace("\\", "/")
    if not normalized:
        raise PaperDocxExportError("manuscript_relative_path is required")
    if normalized.startswith("/") or ".." in normalized.split("/"):
        raise PaperDocxExportError("manuscript_relative_path must stay inside project")
    if not normalized.startswith("manuscript/") or not normalized.endswith(".md"):
        raise PaperDocxExportError("manuscript_relative_path must be a Markdown file under manuscript/")
    return normalized


def _scan_text(relative_path: str, text: str) -> list[str]:
    warnings: list[str] = []
    for pattern in SECRET_PATTERNS:
        if pattern.search(text):
            warnings.append(f"{relative_path}: secret-like pattern redacted from DOCX draft")
            break
    for pattern in ABSOLUTE_PATH_PATTERNS:
        if pattern.search(text):
            warnings.append(f"{relative_path}: absolute-path-like pattern redacted from DOCX draft")
            break
    return warnings


def _sanitize_text(text: str) -> str:
    cleaned = text
    for pattern in SECRET_PATTERNS:
        cleaned = pattern.sub("[redacted-secret]", cleaned)
    for pattern in ABSOLUTE_PATH_PATTERNS:
        cleaned = pattern.sub("[redacted-local-path]", cleaned)
    return cleaned


def _artifact_entry(
    project_dir: Path,
    relative_path: str,
    artifact_type: str,
    mime_type: str,
    required: bool = True,
) -> dict[str, Any]:
    path = _project_path(project_dir, relative_path)
    if not path.exists():
        return {
            "artifact_type": artifact_type,
            "relative_path": relative_path,
            "mime_type": mime_type,
            "available": False,
            "required": required,
            "size_bytes": 0,
            "sha256": None,
        }
    return {
        "artifact_type": artifact_type,
        "relative_path": relative_path,
        "mime_type": mime_type,
        "available": True,
        "required": required,
        "size_bytes": path.stat().st_size,
        "sha256": _sha256(path),
    }


def _add_markdown_line(document: Document, line: str) -> None:
    cleaned = _sanitize_text(line.strip())
    if not cleaned:
        return
    if cleaned.startswith("# "):
        document.add_heading(cleaned[2:].strip(), level=1)
        return
    if cleaned.startswith("## "):
        document.add_heading(cleaned[3:].strip(), level=2)
        return
    if cleaned.startswith("### "):
        document.add_heading(cleaned[4:].strip(), level=3)
        return
    if cleaned.startswith("- "):
        document.add_paragraph(cleaned[2:].strip(), style="List Bullet")
        return
    if cleaned.startswith(">"):
        paragraph = document.add_paragraph()
        run = paragraph.add_run(cleaned.lstrip("> ").strip())
        run.italic = True
        return
    document.add_paragraph(cleaned)


def _write_docx(
    project_dir: Path,
    project_id: str,
    source_markdown_file: str,
    docx_file: str,
    title: str,
) -> list[str]:
    source_text = _read_source_markdown(project_dir, source_markdown_file)
    warnings = _scan_text(source_markdown_file, source_text)
    generated_at = _utc_now()
    document = Document()
    document.add_heading(title, level=0)
    document.add_paragraph(f"Project ID: {project_id}")
    document.add_paragraph(f"Source markdown: {source_markdown_file}")
    document.add_paragraph(f"Generated at: {generated_at}")
    caveat = document.add_paragraph()
    run = caveat.add_run(DOCX_DRAFT_CAVEAT)
    run.bold = True
    document.add_paragraph(
        "This export is a review convenience artifact. It must not be treated as citation proof, "
        "peer review, scientific validation, or publication readiness."
    )
    for line in source_text.splitlines():
        _add_markdown_line(document, line)
    target_path = _project_path(project_dir, docx_file)
    ensure_dir(target_path.parent)
    document.save(target_path)
    return warnings


def _write_manifest(
    project_dir: Path,
    project_id: str,
    source_markdown_file: str,
    docx_file: str,
    manifest_file: str,
    warnings: list[str],
    export_kind: str,
) -> dict[str, Any]:
    generated_at = _utc_now()
    payload = {
        "schema_version": "researchagent.paper_docx_export.v1",
        "project_id": project_id,
        "created_at": generated_at,
        "export_kind": export_kind,
        "source_markdown_file": source_markdown_file,
        "docx_file": docx_file,
        "manifest_file": manifest_file,
        "artifact": _artifact_entry(
            project_dir,
            docx_file,
            "docx_draft",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ),
        "source": _artifact_entry(project_dir, source_markdown_file, "source_markdown", "text/markdown"),
        "safety": {
            "project_relative_paths_only": True,
            "secret_scan_passed": not warnings,
            "warning_count": len(warnings),
        },
        "warnings": warnings,
        "caveats": [
            DOCX_DRAFT_CAVEAT,
            "DOCX export does not approve references, verify citations, or upgrade evidence trust status.",
            "Human review is required before external manuscript use.",
        ],
        "is_draft_artifact": True,
        "citation_proof": False,
        "evidence_trust_package_citation_proof": False,
    }
    write_json(_project_path(project_dir, manifest_file), payload)
    return payload


def read_docx_export_status(
    project_dir: Path,
    *,
    docx_file: str = DRAFT_FULL_DOCX,
    manifest_file: str = DRAFT_FULL_DOCX_MANIFEST,
) -> dict[str, Any]:
    path = _project_path(project_dir, docx_file)
    manifest_path = _project_path(project_dir, manifest_file)
    return {
        "available": path.exists(),
        "docx_file": docx_file if path.exists() else None,
        "manifest_file": manifest_file if manifest_path.exists() else None,
        "size_bytes": path.stat().st_size if path.exists() else 0,
        "sha256": _sha256(path) if path.exists() else None,
    }


def export_draft_docx(project_dir: Path, project_id: str) -> dict[str, Any]:
    warnings = _write_docx(
        project_dir,
        project_id,
        DRAFT_FULL_MD,
        DRAFT_FULL_DOCX,
        "ResearchAgent Auto Paper Writer Draft",
    )
    manifest = _write_manifest(
        project_dir,
        project_id,
        DRAFT_FULL_MD,
        DRAFT_FULL_DOCX,
        DRAFT_FULL_DOCX_MANIFEST,
        warnings,
        "paper_writer_draft",
    )
    append_audit_event(
        project_dir,
        project_id,
        "export_auto_paper_docx",
        "Auto Paper Writer exported a DOCX draft from Markdown.",
        {"docx_file": DRAFT_FULL_DOCX, "source_markdown_file": DRAFT_FULL_MD, "manifest_file": DRAFT_FULL_DOCX_MANIFEST},
        source="api",
        event_category="manuscript",
        risk_level="low",
        entity_type="manuscript",
        entity_id="draft_full_docx",
    )
    return manifest


def export_auto_scientist_paper_docx(
    project_dir: Path,
    project_id: str,
    manuscript_relative_path: str | None = None,
) -> dict[str, Any]:
    source_markdown_file = _validate_manuscript_markdown_path(manuscript_relative_path or AUTONOMOUS_PAPER_MD)
    warnings = _write_docx(
        project_dir,
        project_id,
        source_markdown_file,
        AUTO_SCIENTIST_PAPER_DOCX,
        "ResearchAgent Auto Scientist Paper Draft",
    )
    manifest = _write_manifest(
        project_dir,
        project_id,
        source_markdown_file,
        AUTO_SCIENTIST_PAPER_DOCX,
        AUTO_SCIENTIST_PAPER_DOCX_MANIFEST,
        warnings,
        "auto_scientist_paper",
    )
    append_audit_event(
        project_dir,
        project_id,
        "export_auto_scientist_paper_docx",
        "Auto Scientist exported a DOCX paper draft from Markdown.",
        {
            "docx_file": AUTO_SCIENTIST_PAPER_DOCX,
            "source_markdown_file": source_markdown_file,
            "manifest_file": AUTO_SCIENTIST_PAPER_DOCX_MANIFEST,
        },
        source="api",
        event_category="auto_scientist",
        risk_level="low",
        entity_type="manuscript",
        entity_id="auto_scientist_paper_docx",
    )
    return manifest
